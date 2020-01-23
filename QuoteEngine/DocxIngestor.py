from .IngestorInterface import IngestorInterface
from .QuoteModel import QuoteModel
from typing import List
import docx

class DocxIngestor(IngestorInterface):

    allowed = ['docx']

    @classmethod
    def parse(cls, path: str) -> List[QuoteModel]:

        if cls.can_ingest(path):


            temp = []

            doc = docx.Document(path)

            data = [b.text for b in doc.paragraphs]

            for d in data:
                parse = d.split('-')
                if len(parse) == 2:
                    new = QuoteModel(parse[0],parse[1])
                    temp.append(new)
            return temp


            '''       
            for para in doc.paragraphs:
                if para.text != "":
                    parse = para.text.split('-')
            
                    new = QuoteModel(parse[0],parse[1])
                    temp.append(new)
            '''
